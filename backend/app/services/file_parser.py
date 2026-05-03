"""
Multi-format File Parser Service

Supports: PDF, DOCX, DOC, XLSX, XLS, TXT, and images (PNG, JPG, etc.)
"""
import os
import io
from typing import Optional, Tuple
from pathlib import Path

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

try:
    from pdf2image import convert_from_path
except ImportError:
    convert_from_path = None

try:
    from docx import Document
except ImportError:
    Document = None

try:
    import openpyxl
    import pandas as pd
except ImportError:
    openpyxl = None
    pd = None

try:
    import pytesseract
    from PIL import Image
except ImportError:
    pytesseract = None
    Image = None


class FileParser:
    """文件解析器"""

    SUPPORTED_FORMATS = {
        "pdf", "doc", "docx", "xls", "xlsx", "txt",
        "png", "jpg", "jpeg", "bmp", "gif"
    }

    @staticmethod
    def parse(file_path: str) -> Tuple[str, Optional[str]]:
        """
        Parse file and extract text content
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (extracted_text, error_message)
        """
        try:
            # Get file extension
            file_ext = Path(file_path).suffix.lower().lstrip(".")

            if file_ext == "pdf":
                return FileParser._parse_pdf(file_path)
            elif file_ext in ["docx"]:
                return FileParser._parse_docx(file_path)
            elif file_ext in ["doc"]:
                return FileParser._parse_doc(file_path)
            elif file_ext in ["xlsx"]:
                return FileParser._parse_xlsx(file_path)
            elif file_ext in ["xls"]:
                return FileParser._parse_xls(file_path)
            elif file_ext == "txt":
                return FileParser._parse_txt(file_path)
            elif file_ext in ["png", "jpg", "jpeg", "bmp", "gif"]:
                return FileParser._parse_image(file_path)
            else:
                return "", f"Unsupported file format: {file_ext}"
                
        except Exception as e:
            return "", f"Error parsing file: {str(e)}"

    @staticmethod
    def _parse_pdf(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse PDF file"""
        try:
            if not pdfplumber:
                return "", "pdfplumber not installed"
                
            text_content = []
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    text = page.extract_text() or ""
                    text_content.append(f"\n--- Page {page_num} ---\n{text}")
                    
                    # Extract tables if any
                    tables = page.extract_tables()
                    if tables:
                        for table in tables:
                            table_text = "\n".join(
                                [" | ".join([str(cell) if cell else "" for cell in row]) for row in table]
                            )
                            text_content.append(f"\n[TABLE]\n{table_text}\n[/TABLE]\n")
            
            return "\n".join(text_content), None
        except Exception as e:
            return "", f"Error parsing PDF: {str(e)}"

    @staticmethod
    def _parse_docx(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse DOCX file"""
        try:
            if not Document:
                return "", "python-docx not installed"
                
            doc = Document(file_path)
            text_content = []
            
            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
            
            # Extract tables
            for table in doc.tables:
                table_text = []
                for row in table.rows:
                    row_text = [cell.text for cell in row.cells]
                    table_text.append(" | ".join(row_text))
                text_content.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")
            
            return "\n".join(text_content), None
        except Exception as e:
            return "", f"Error parsing DOCX: {str(e)}"

    @staticmethod
    def _parse_doc(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse DOC file (try to convert to DOCX first)"""
        # Try to parse as DOCX first (many .doc files are actually DOCX format)
        text, error = FileParser._parse_docx(file_path)
        if not error:
            return text, None
        return "", "DOC format requires additional libraries (python-docx may not support old .doc)"

    @staticmethod
    def _parse_xlsx(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse XLSX file"""
        try:
            if not pd:
                return "", "pandas not installed"
                
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                text_content.append(f"\n--- Sheet: {sheet_name} ---\n")
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                
                # Convert dataframe to text
                text_content.append(df.to_string())
            
            return "\n".join(text_content), None
        except Exception as e:
            return "", f"Error parsing XLSX: {str(e)}"

    @staticmethod
    def _parse_xls(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse XLS file"""
        try:
            if not pd:
                return "", "pandas not installed"
                
            excel_file = pd.ExcelFile(file_path)
            text_content = []
            
            for sheet_name in excel_file.sheet_names:
                text_content.append(f"\n--- Sheet: {sheet_name} ---\n")
                df = pd.read_excel(file_path, sheet_name=sheet_name)
                text_content.append(df.to_string())
            
            return "\n".join(text_content), None
        except Exception as e:
            return "", f"Error parsing XLS: {str(e)}"

    @staticmethod
    def _parse_txt(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse TXT file"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            return content, None
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, "r", encoding="gbk") as f:
                    content = f.read()
                return content, None
            except Exception as e:
                return "", f"Error parsing TXT: {str(e)}"
        except Exception as e:
            return "", f"Error parsing TXT: {str(e)}"

    @staticmethod
    def _parse_image(file_path: str) -> Tuple[str, Optional[str]]:
        """Parse image file using OCR"""
        try:
            if not pytesseract or not Image:
                return "", "pytesseract or Pillow not installed"
                
            # Check if Tesseract is installed
            try:
                pytesseract.get_tesseract_version()
            except pytesseract.TesseractNotFoundError:
                return "", "Tesseract OCR not installed on system"
            
            # Open image and perform OCR
            image = Image.open(file_path)
            text = pytesseract.image_to_string(image, lang="chi_sim+eng")
            return text, None
            
        except Exception as e:
            return "", f"Error parsing image: {str(e)}"

    @staticmethod
    def get_file_size(file_path: str) -> int:
        """Get file size in bytes"""
        try:
            return os.path.getsize(file_path)
        except Exception:
            return 0
